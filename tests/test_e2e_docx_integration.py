"""Test de integracion: build_book_docx como capability registrada.

Verifica que:
- CONFIG["docx"] == True en el runner E2E.
- La capability 'build_book_docx' esta registrada en core/schemas.
- build_book_docx produce un DOCX valido (existe, .docx, size>0, abre con python-docx).
- La prioridad edited_es > draft_es se respeta en el libro construido.
"""


from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path

import pytest

from core.book.book_schema import Book, Chapter
from core.schemas import PAYLOAD_SCHEMAS
from modules.document_builder.main import build_book_docx


def test_config_docx_is_true():
    """El runner E2E debe tener CONFIG con docx=True."""
    import importlib
    import run_e2e_001_editorial as runner
    importlib.reload(runner)
    assert runner.CONFIG["docx"] is True, "CONFIG['docx'] debe ser True para generar DOCX en el E2E"


def test_build_book_docx_capability_registered():
    """La capability build_book_docx debe estar registrada en core/schemas."""
    assert "build_book_docx" in PAYLOAD_SCHEMAS, (
        "build_book_docx debe estar registrada como capability"
    )


def test_build_book_docx_produces_valid_docx(tmp_path: Path) -> None:
    """build_book_docx debe producir un DOCX que cumple todos los requisitos."""
    # Construimos un book como el E2E lo prepara
    book = Book(
        book_id=1001,
        title="El nacimiento de Internet",
        subtitle="Desde ARPANET hasta TCP/IP",
        description="Introducción al libro sobre Internet.",
        author="Space Lair",
        target_audience="General",
        genre="tecnologia",
        languages=["es"],
        target_chapters=1,
        status="edited",
        created_at=datetime(2024, 1, 1),
        chapters=[
            Chapter(
                chapter_id=1,
                book_id=1001,
                number=1,
                title="El nacimiento de Internet",
                edited_es="## Introducción\n\nContenido editado del capítulo 1.",
                draft_es="## Borrador\n\nContenido del borrador.",
                images=[],
            )
        ],
    )
    payload = {
        "book": book.model_dump(),
        "language": "es",
        "page_config": {"size": "A4", "margins_mm": {"top": 25.4, "bottom": 25.4, "left": 25.4, "right": 25.4}},
    }
    out = build_book_docx(payload)

    docx_path = out["docx_path"]

    # Verificaciones
    assert os.path.isfile(docx_path), "El archivo DOCX debe existir"
    assert docx_path.endswith(".docx"), "La extensión debe ser .docx"
    assert os.path.getsize(docx_path) > 0, "El archivo DOCX no debe estar vacío"

    from docx import Document
    doc = Document(docx_path)
    assert doc is not None, "python-docx debe poder abrir el archivo"

    assert out["book_id"] == 1001
    assert out["language"] == "es"
    assert out["chapter_count"] == 1
    assert out["image_count"] == 0


def test_build_book_docx_prefers_edited_over_draft(tmp_path: Path) -> None:
    """El contenido edited_es debe aparecer en el DOCX, no draft_es."""
    book = Book(
        book_id=1001,
        title="Libro test",
        author="Space Lair",
        genre="tecnologia",
        languages=["es"],
        target_chapters=1,
        status="edited",
        chapters=[
            Chapter(
                chapter_id=1,
                book_id=1001,
                number=1,
                title="Capítulo 1",
                edited_es="## Editado\n\nContenido EDITADO priority.",
                draft_es="## Borrador\n\nContenido BORRADOR que NO debe aparecer.",
                images=[],
            )
        ],
    )
    payload = {"book": book.model_dump(), "language": "es", "page_config": None}
    out = build_book_docx(payload)

    from docx import Document
    doc = Document(out["docx_path"])
    texts = [p.text for p in doc.paragraphs]
    assert any("Contenido EDITADO priority" in t for t in texts), "Debe aparecer edited_es"
    assert not any("Contenido BORRADOR que NO debe aparecer" in t for t in texts), "draft_es NO debe aparecer"



def test_build_book_docx_long_description_does_not_crash(tmp_path: Path) -> None:
    """Una description >255 chars no debe provocar ValueError en core_properties.comments.

    Regresion: el runner anteriormente pasaba ctext[:500] como description,
    lo que provocaba 'exceeded 255 char limit' en el Document Builder.
    El guard [:255] debe truncar sin lanzar excepcion.
    """
    long_desc = "X" * 500  # 500 chars, supera el limite de 255 de DOCX
    book = Book(
        book_id=1001,
        title="Libro test",
        description=long_desc,
        author="Space Lair",
        genre="tecnologia",
        languages=["es"],
        target_chapters=1,
        status="edited",
        chapters=[
            Chapter(
                chapter_id=1,
                book_id=1001,
                number=1,
                title="Capitulo 1",
                edited_es="## Introduccion\n\nContenido editado.",
                draft_es="Borrador",
                images=[],
            )
        ],
    )
    payload = {"book": book.model_dump(), "language": "es", "page_config": None}
    # No debe lanzar ValueError: exceeded 255 char limit
    out = build_book_docx(payload)
    assert os.path.isfile(out["docx_path"])
    assert os.path.getsize(out["docx_path"]) > 0

    # Verificar: el comentario del DOCX esta truncado a 255
    from docx import Document
    doc = Document(out["docx_path"])
    comments = doc.core_properties.comments or ""
    assert len(comments) <= 255
