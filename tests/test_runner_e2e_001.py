"""Tests específicos del runner E2E 001."""

from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import patch

import pytest

import run_e2e_001_editorial as runner
from run_e2e_001_editorial import _json_safe, _safe_book_description, print_report


def test_fallback_chapter_does_not_reach_completed(monkeypatch, tmp_path: Path) -> None:
    """Capítulo FAIL corta el pipeline en chapter; nunca llega a QC ni completed."""
    import importlib

    importlib.reload(runner)

    monkeypatch.setattr(runner, "init_db", lambda: None)
    monkeypatch.setattr(runner, "plan_execute", lambda p: {
        "title": "Test", "provider": "ollama", "model": "test",
        "chapters": [{"image_requirements": 0}],
    })
    monkeypatch.setattr(runner, "research_execute", lambda p: {
        "status": "PASS",
        "sources": [
            {
                "title": f"s{i}",
                "url": f"https://space-lair.test/{i}",
                "source_type": "web",
                "relevance": "N/A",
                "accessed_at": "2024-01-01T00:00:00+00:00",
            }
            for i in range(3)
        ],
        "execution_mode": "real",
    })
    monkeypatch.setattr(runner, "chapter_execute", lambda p: {
        "word_count": 871,
        "quality_gate": "FAIL",
        "execution_mode": "real",
        "quality_errors": ["menos de 1500 palabras (871)"],
        "chapter_md_path": "dummy.md",
    })
    monkeypatch.setattr(runner, "_chapter_detect_placeholder", lambda t: False)
    monkeypatch.setattr(runner, "save_checkpoint", lambda *a, **k: "cp.json")

    result = runner.run()

    assert result["status"] == "error"
    assert result["failed_stage"] == "chapter"
    assert result["chapter_generation_status"] == "FAIL"
    assert result["chapter_generation_word_count"] == 871
    assert "menos de 1500 palabras (871)" in (result.get("error") or "")


def test_quality_gate_fails_when_under_minimum_words(monkeypatch, tmp_path: Path) -> None:
    """Capítulo efectivo de 1200 palabras → Quality Gate FAIL."""
    import importlib

    importlib.reload(runner)

    monkeypatch.setattr(runner, "init_db", lambda: None)
    monkeypatch.setattr(runner, "plan_execute", lambda p: {
        "title": "Test", "provider": "ollama", "model": "test",
        "chapters": [{"image_requirements": 0}],
    })
    monkeypatch.setattr(runner, "research_execute", lambda p: {
        "status": "PASS",
        "sources": [
            {
                "title": f"s{i}",
                "url": f"https://space-lair.test/{i}",
                "source_type": "web",
                "relevance": "N/A",
                "accessed_at": "2024-01-01T00:00:00+00:00",
            }
            for i in range(3)
        ],
        "execution_mode": "real",
    })
    monkeypatch.setattr(runner, "chapter_execute", lambda p: {
        "word_count": 1200,
        "quality_gate": "FAIL",
        "execution_mode": "real",
        "quality_errors": ["menos de 1500 palabras (1200)"],
        "chapter_md_path": "dummy.md",
    })
    monkeypatch.setattr(runner, "_chapter_detect_placeholder", lambda t: False)
    from docx import Document

    def mock_build_2(p):
        docx_path = str(tmp_path / "book.docx")
        doc = Document()
        doc.add_paragraph("Índice")
        doc.add_paragraph("Capítulo 1: El nacimiento de Internet")
        doc.add_paragraph("palabra " * 1200)
        doc.save(docx_path)
        return {
            "docx_path": docx_path,
            "book_id": 1001,
            "language": "es",
            "chapter_count": 1,
            "image_count": 0,
        }

    monkeypatch.setattr(runner, "build_book_docx", mock_build_2)
    monkeypatch.setattr(runner, "save_checkpoint", lambda *a, **k: "cp.json")

    result = runner.run()

    # Chapter Writer FAIL corta el pipeline antes del QC.
    assert result["status"] == "error"
    assert result["failed_stage"] == "chapter"
    assert result["chapter_generation_status"] == "FAIL"
    assert result["chapter_generation_word_count"] == 1200
    assert "menos de 1500 palabras (1200)" in (result.get("error") or "")


def test_valid_generation_1800_passes(monkeypatch, tmp_path: Path) -> None:
    """Generación válida de 1800+ → Quality Gate PASS."""
    import importlib

    importlib.reload(runner)

    monkeypatch.setattr(runner, "init_db", lambda: None)
    monkeypatch.setattr(runner, "plan_execute", lambda p: {
        "title": "Test", "provider": "ollama", "model": "test",
        "chapters": [{"image_requirements": 0}],
    })
    monkeypatch.setattr(runner, "research_execute", lambda p: {
        "status": "PASS",
        "sources": [
            {
                "title": f"s{i}",
                "url": f"https://space-lair.test/{i}",
                "source_type": "web",
                "relevance": "N/A",
                "accessed_at": "2024-01-01T00:00:00+00:00",
            }
            for i in range(3)
        ],
        "execution_mode": "real",
    })
    chapter_file = tmp_path / "chapter.md"
    chapter_file.write_text("palabra " * 1800, encoding="utf-8")
    monkeypatch.setattr(runner, "chapter_execute", lambda p: {
        "word_count": 1800,
        "quality_gate": "PASS",
        "execution_mode": "real",
        "quality_errors": [],
        "chapter_md_path": str(chapter_file),
    })
    monkeypatch.setattr(runner, "_chapter_detect_placeholder", lambda t: False)
    monkeypatch.setattr(runner, "fc_execute", lambda p: {
        "status": "PASS",
        "quality_gate": "PASS",
        "claims_checked": 1,
        "supported_claims": 0,
        "unsupported_claims": 0,
        "conflicting_claims": 0,
    })
    monkeypatch.setattr(runner, "editor_execute", lambda p: {
        "status": "PASS",
        "quality_gate": "PASS",
        "input_words": 1800,
        "output_words": 1800,
        "execution_mode": "real",
        "placeholder_detected": False,
        "edited_text": "palabra " * 1800,
    })
    from docx import Document

    def mock_build_3(p):
        docx_path = str(tmp_path / "book.docx")
        doc = Document()
        doc.add_paragraph("Índice")
        doc.add_paragraph("Capítulo 1: El nacimiento de Internet")
        doc.add_paragraph("palabra " * 1800)
        doc.save(docx_path)
        return {
            "docx_path": docx_path,
            "book_id": 1001,
            "language": "es",
            "chapter_count": 1,
            "image_count": 0,
        }

    monkeypatch.setattr(runner, "build_book_docx", mock_build_3)
    monkeypatch.setattr(runner, "save_checkpoint", lambda *a, **k: "cp.json")

    result = runner.run()

    assert result["quality_gate"] == "PASS"
    assert result["chapter_word_count"] == 1800
    assert result["effective_chapter_word_count"] == 1800
    assert result["chapter_status"] == "PASS"
    assert result["status"] == "completed"


def test_json_safe_no_circular_reference() -> None:
    """Reporte sin referencias circulares produce JSON válido."""
    data = {"a": 1, "b": [2, 3], "c": None, "d": "text"}
    safe = _json_safe(data)
    dumped = json.dumps(safe, ensure_ascii=False)
    assert json.loads(dumped) == data


def test_json_safe_circular_reference_marker() -> None:
    """Payload con referencia circular → reporte no falla y contiene marcador."""
    payload: dict = {"x": 1}
    payload["self"] = payload
    safe = _json_safe(payload)
    dumped = json.dumps(safe, ensure_ascii=False)
    assert "<circular_reference>" in dumped
    assert safe["self"] == "<circular_reference>"


def test_print_report_no_traceback_on_circular_payload() -> None:
    """print_report no debe producir traceback con referencias circulares."""
    payload: dict = {"a": 1}
    payload["self"] = payload
    report = {
        "status": "error",
        "failed_stage": "quality_gate",
        "error": "boom",
        "traceback": "Traceback...",
        "payload": payload,
    }
    # No debe lanzar excepción
    print_report(report)


def test_safe_book_description_short_text_preserved() -> None:
    """Descripción corta se conserva sin cambios."""
    original = "Una historia breve sobre tecnología."
    result = _safe_book_description(original, "Título")
    assert result == original


def test_safe_book_description_long_text_truncated() -> None:
    """Descripción >255 queda <=200."""
    original = "a" * 300
    result = _safe_book_description(original, "Título")
    assert len(result) <= 200
    assert result.endswith("...")


def test_safe_book_description_chapter_content_replaced() -> None:
    """Si parece contenido de capítulo, se reemplaza por descripción corta."""
    chapter_text = "# Introducción\n\n" + ("palabra " * 200)
    result = _safe_book_description(chapter_text, "El nacimiento de Internet")
    assert result == "Libro sobre El nacimiento de Internet."
    assert len(result) <= 200


def test_safe_book_description_chapter_text_not_modified() -> None:
    """La función no modifica el texto original del capítulo."""
    chapter_text = "# Introducción\n\nContenido del capítulo."
    original = chapter_text
    _safe_book_description(chapter_text, "Título")
    assert chapter_text == original


# ============================================
# Coherencia pipeline: texto editado en el DOCX
# ============================================

def _run_editorial_harness(monkeypatch, tmp_path: Path, *, edited_text, chapter_ctext,
                           build_capture: list) -> dict:
    """Ejecuta runner.run() capturando el book_dict pasado a build_book_docx."""
    import importlib
    importlib.reload(runner)

    chapter_file = tmp_path / "chapter.md"
    chapter_file.write_text(chapter_ctext, encoding="utf-8")

    real_open = open

    def mock_open(path, *args, **kwargs):
        path_str = str(path)
        if path_str.endswith("chapter.md"):
            return real_open(chapter_file, *args, **kwargs)
        return real_open(path, *args, **kwargs)

    monkeypatch.setattr("builtins.open", mock_open)
    monkeypatch.setattr(runner, "init_db", lambda: None)
    monkeypatch.setattr(runner, "plan_execute", lambda p: {
        "title": "Test", "provider": "ollama", "model": "test",
        "chapters": [{"image_requirements": 0}],
    })
    monkeypatch.setattr(runner, "research_execute", lambda p: {
        "status": "PASS",
        "sources": [
            {"title": f"s{i}", "url": f"https://space-lair.test/{i}",
             "source_type": "web", "relevance": "N/A",
             "accessed_at": "2024-01-01T00:00:00+00:00"}
            for i in range(3)
        ],
        "execution_mode": "real",
    })
    monkeypatch.setattr(runner, "chapter_execute", lambda p: {
        "word_count": len(chapter_ctext.split()),
        "quality_gate": "PASS",
        "execution_mode": "real",
        "quality_errors": [],
        "chapter_md_path": str(chapter_file),
    })
    monkeypatch.setattr(runner, "_chapter_detect_placeholder", lambda t: False)
    monkeypatch.setattr(runner, "fc_execute", lambda p: {
        "status": "PASS", "quality_gate": "PASS",
        "claims_checked": 1, "supported_claims": 0,
        "unsupported_claims": 0, "conflicting_claims": 0,
    })

    def fake_editor(p):
        out = {"status": "PASS", "quality_gate": "PASS",
               "input_words": len(p.get("chapter_text", "").split()),
               "execution_mode": "real", "placeholder_detected": False}
        if edited_text is not None:
            out["edited_text"] = edited_text
            out["output_words"] = len(edited_text.split())
        else:
            out["output_words"] = len(p.get("chapter_text", "").split())
        return out

    monkeypatch.setattr(runner, "editor_execute", fake_editor)

    def mock_build(p):
        from docx import Document
        build_capture.append(p)
        docx_path = str(tmp_path / "book.docx")
        doc = Document()
        doc.add_paragraph("Índice")
        doc.add_paragraph("Capítulo 1: El nacimiento de Internet")
        doc.add_paragraph("palabra " * 1800)
        doc.save(docx_path)
        return {"docx_path": docx_path, "book_id": 1001,
                "language": "es", "chapter_count": 1, "image_count": 0}

    monkeypatch.setattr(runner, "build_book_docx", mock_build)
    monkeypatch.setattr(runner, "save_checkpoint", lambda *a, **k: "cp.json")

    return runner.run()


def test_docx_uses_editor_text_when_different(monkeypatch, tmp_path: Path) -> None:
    """Si el Editor devuelve texto distinto de ctext, el DOCX usa edited_es editado."""
    from docx import Document  # noqa: F401
    ctext = "palabra " * 1800
    edited = "TEXTO EDITADO POR EL EDITOR " + ("palabra " * 1800)
    capture: list = []
    result = _run_editorial_harness(
        monkeypatch, tmp_path, edited_text=edited, chapter_ctext=ctext, build_capture=capture)
    assert result["status"] == "completed"
    assert len(capture) == 1
    book_dict = capture[0]["book"]
    assert book_dict["chapters"][0]["edited_es"] == edited
    assert book_dict["chapters"][0]["edited_es"] != ctext


def test_docx_falls_back_to_ctext_when_no_editor_text(monkeypatch, tmp_path: Path) -> None:
    """Si el Editor no devuelve edited_text, se usa ctext como fallback."""
    from docx import Document  # noqa: F401
    ctext = "palabra " * 1800
    capture: list = []
    result = _run_editorial_harness(
        monkeypatch, tmp_path, edited_text=None, chapter_ctext=ctext, build_capture=capture)
    assert result["status"] == "completed"
    assert len(capture) == 1
    book_dict = capture[0]["book"]
    assert book_dict["chapters"][0]["edited_es"] == ctext


def test_docx_description_not_from_chapter_body(monkeypatch, tmp_path: Path) -> None:
    """book.description no contiene ctext[:500] y es <=200 caracteres."""
    from docx import Document  # noqa: F401
    ctext = "# Introducción al capítulo\n\n" + ("palabra " * 1800)
    edited = "palabra " * 1800
    capture: list = []
    result = _run_editorial_harness(
        monkeypatch, tmp_path, edited_text=edited, chapter_ctext=ctext, build_capture=capture)
    assert result["status"] == "completed"
    desc = capture[0]["book"]["description"] or ""
    assert isinstance(desc, str)
    assert len(desc) <= 200
    # No debe contener el arranque crudo del capítulo
    assert "# Introducción" not in desc
    assert "## " not in desc


def test_docx_description_shorter_than_200(monkeypatch, tmp_path: Path) -> None:
    """Descripción del DOCX siempre <= 200 caracteres."""
    from docx import Document  # noqa: F401
    ctext = "palabra " * 1800
    capture: list = []
    _run_editorial_harness(
        monkeypatch, tmp_path, edited_text="palabra " * 1800,
        chapter_ctext=ctext, build_capture=capture)
    desc = capture[0]["book"]["description"] or ""
    assert len(desc) <= 200


def test_build_book_docx_includes_edited_text(tmp_path: Path) -> None:
    """build_book_docx real pone edited_es en el DOCX (no draft_es)."""
    from docx import Document
    from modules.document_builder.main import build_book_docx

    edited_text = ("## Editado final\n\n" + ("palabra " * 400))
    book_dict = {
        "book_id": 1001,
        "title": "El nacimiento de Internet",
        "description": "Libro sobre El nacimiento de Internet.",
        "author": "Space Lair",
        "languages": ["es"],
        "target_chapters": 1,
        "status": "edited",
        "chapters": [
            {
                "chapter_id": 1,
                "book_id": 1001,
                "number": 1,
                "title": "Capítulo 1",
                "edited_es": edited_text,
                "draft_es": "BORRADOR ORIGINAL DEL CHAPTER WRITER",
                "images": [],
            }
        ],
    }
    out = build_book_docx({"book": book_dict, "language": "es", "page_config": None})
    doc = Document(out["docx_path"])
    texts = [p.text for p in doc.paragraphs]
    assert any("Editado final" in t for t in texts)
    assert not any("BORRADOR ORIGINAL DEL CHAPTER WRITER" in t for t in texts)



def test_effective_word_count_uses_editor_output(monkeypatch, tmp_path: Path) -> None:
    """effective_chapter_word_count deriva del texto EDITADO, no del Chapter Writer.

    El DOCX contiene edited_es (texto del editor). Si el Editor acorta el texto,
    el word count efectivo verificado por el QC debe ser el del texto editado.
    """
    from docx import Document  # noqa: F401
    # Chapter Writer produce 2000 palabras (>= minimum_words 1500, PASS)
    ctext = "palabra " * 2000
    # Editor produce 1800 palabras (texto editado mas corto, pero >= 1500)
    edited = "palabra " * 1800
    capture: list = []
    result = _run_editorial_harness(
        monkeypatch, tmp_path, edited_text=edited, chapter_ctext=ctext,
        build_capture=capture)
    assert result["status"] == "completed"
    assert result["quality_gate"] == "PASS"
    # effective_chapter_word_count debe ser del EDITOR (1800), no del capitulo (2000)
    assert result["effective_chapter_word_count"] == 1800
    assert result["chapter_generation_word_count"] == 2000
    assert result["effective_chapter_word_count"] != result["chapter_generation_word_count"]
    # El DOCX contiene el texto editado, no el original
    assert capture[0]["book"]["chapters"][0]["edited_es"] == edited


def test_quality_gate_fails_when_editor_shortens_below_minimum(monkeypatch, tmp_path: Path) -> None:
    """Si el Editor acorta el texto por debajo de minimum_words, el runner falla en editor.

    El runner valida editor_output_words < minimum_words y corta antes del QC.
    """
    from docx import Document  # noqa: F401
    ctext = "palabra " * 2000  # 2000 words, >= minimum_words, chapter PASS
    edited = "palabra " * 1400  # 1400 words, < minimum_words (1500)
    capture: list = []
    result = _run_editorial_harness(
        monkeypatch, tmp_path, edited_text=edited, chapter_ctext=ctext,
        build_capture=capture)
    assert result["status"] == "error"
    assert result["failed_stage"] == "editor"
    # effective_chapter_word_count se actualiza DESPUÉS del editor, pero el runner
    # corta en editor antes de hacerlo; por tanto, mantiene el valor del chapter.
    assert result["effective_chapter_word_count"] == 2000
    assert result["chapter_generation_word_count"] == 2000
    assert "editor output words (1400) < minimum_words (1500)" in (result.get("error") or "")
    # build_book_docx nunca se llama porque el runner falló en editor
    assert len(capture) == 0


def test_final_checkpoint_contains_editor_metadata(monkeypatch, tmp_path: Path) -> None:
    """El checkpoint final_final_qc debe contener los campos del editor y coincidir con el report."""

    ctext = "palabra " * 2500
    edited = "palabra " * 2500
    capture: list = []
    saved_checkpoints: list = []

    def _run():
        import importlib
        importlib.reload(runner)

        chapter_file = tmp_path / "chapter.md"
        chapter_file.write_text(ctext, encoding="utf-8")

        real_open = open

        def mock_open(path, *args, **kwargs):
            path_str = str(path)
            if path_str.endswith("chapter.md"):
                return real_open(chapter_file, *args, **kwargs)
            return real_open(path, *args, **kwargs)

        monkeypatch.setattr("builtins.open", mock_open)
        monkeypatch.setattr(runner, "init_db", lambda: None)
        monkeypatch.setattr(runner, "plan_execute", lambda p: {
            "title": "Test", "provider": "ollama", "model": "test",
            "chapters": [{"image_requirements": 0}],
        })
        monkeypatch.setattr(runner, "research_execute", lambda p: {
            "status": "PASS",
            "sources": [
                {
                    "title": f"s{i}", "url": f"https://space-lair.test/{i}",
                    "source_type": "web", "relevance": "N/A",
                    "accessed_at": "2024-01-01T00:00:00+00:00",
                }
                for i in range(3)
            ],
            "execution_mode": "real",
        })
        monkeypatch.setattr(runner, "chapter_execute", lambda p: {
            "word_count": len(ctext.split()),
            "quality_gate": "PASS",
            "execution_mode": "real",
            "quality_errors": [],
            "chapter_md_path": str(chapter_file),
        })
        monkeypatch.setattr(runner, "_chapter_detect_placeholder", lambda t: False)
        monkeypatch.setattr(runner, "fc_execute", lambda p: {
            "status": "PASS", "quality_gate": "PASS",
            "claims_checked": 1, "supported_claims": 0,
            "unsupported_claims": 0, "conflicting_claims": 0,
        })

        def fake_editor(p):
            out = {
                "status": "PASS", "quality_gate": "PASS",
                "input_words": len(p.get("chapter_text", "").split()),
                "execution_mode": "real", "placeholder_detected": False,
            }
            if edited is not None:
                out["edited_text"] = edited
                out["output_words"] = len(edited.split())
            else:
                out["output_words"] = len(p.get("chapter_text", "").split())
            return out

        monkeypatch.setattr(runner, "editor_execute", fake_editor)

        def mock_build(p):
            from docx import Document
            capture.append(p)
            docx_path = str(tmp_path / "book.docx")
            doc = Document()
            doc.add_paragraph("Índice")
            doc.add_paragraph("Capítulo 1: El nacimiento de Internet")
            doc.add_paragraph("palabra " * 1800)
            doc.save(docx_path)
            return {"docx_path": docx_path, "book_id": 1001,
                    "language": "es", "chapter_count": 1, "image_count": 0}

        monkeypatch.setattr(runner, "build_book_docx", mock_build)
        monkeypatch.setattr(runner, "save_checkpoint", lambda *a, **k: saved_checkpoints.append((a, k)))
        return runner.run()

    result = _run()
    assert result["status"] == "completed"
    assert result["quality_gate"] == "PASS"

    # El último checkpoint guardado es el final_qc
    assert saved_checkpoints, "save_checkpoint no fue llamado"
    final_args, final_kwargs = saved_checkpoints[-1]
    payload = final_args[2]

    assert payload.get("editor_execution_mode") == "real"
    assert payload.get("editor_placeholder_detected") is False
    assert payload.get("editor_input_words") == 2500
    assert payload.get("editor_output_words") == 2500
    assert payload.get("effective_chapter_word_count") == 2500
    assert payload.get("quality_gate_status") == "PASS"
    assert payload.get("fact_check_status") == "PASS"

    # Los valores deben coincidir con el report final
    assert payload.get("effective_chapter_word_count") == result["effective_chapter_word_count"]
    assert payload.get("editor_output_words") == result.get("editor_output_words", 2500)


def _fallback_setup(monkeypatch, tmp_path: Path, fallback_words: int = 1742,
                    fail_words: int = 871, chapter_gate: str = "FAIL") -> dict:
    """Configura mocks para inducir fallback de capítulo. Devuelve helpers para asserts."""
    import importlib
    importlib.reload(runner)

    fallback_text = "palabra " * fallback_words
    fallback_file = tmp_path / "chapter.md"
    fallback_file.write_text(fallback_text, encoding="utf-8")

    real_open = open

    def mock_open(path, *args, **kwargs):
        if "chapter.md" in str(path):
            return real_open(fallback_file, *args, **kwargs)
        return real_open(path, *args, **kwargs)

    monkeypatch.setattr("builtins.open", mock_open)
    monkeypatch.setattr(runner, "init_db", lambda: None)
    monkeypatch.setattr(runner, "plan_execute", lambda p: {
        "title": "Test", "provider": "ollama", "model": "test",
        "chapters": [{"image_requirements": 0}],
    })
    monkeypatch.setattr(runner, "research_execute", lambda p: {
        "status": "PASS",
        "sources": [
            {"title": f"s{i}", "url": f"https://space-lair.test/{i}",
             "source_type": "web", "relevance": "N/A",
             "accessed_at": "2024-01-01T00:00:00+00:00"}
            for i in range(3)
        ],
        "execution_mode": "real",
    })
    monkeypatch.setattr(runner, "chapter_execute", lambda p: {
        "word_count": fail_words,
        "quality_gate": chapter_gate,
        "execution_mode": "real",
        "quality_errors": [f"menos de 1500 palabras ({fail_words})"],
        "chapter_md_path": "dummy.md",
    })
    monkeypatch.setattr(runner, "_chapter_detect_placeholder", lambda t: False)
    fc_capture: list = []

    def _fc(p):
        fc_capture.append(p.get("chapter_text", ""))
        return {
            "status": "PASS", "quality_gate": "PASS",
            "claims_checked": 1, "supported_claims": 0,
            "unsupported_claims": 0, "conflicting_claims": 0,
        }

    monkeypatch.setattr(runner, "fc_execute", _fc)
    monkeypatch.setattr(runner, "editor_execute", lambda p: {
        "status": "PASS", "quality_gate": "PASS",
        "input_words": fallback_words, "output_words": fallback_words,
        "execution_mode": "fallback", "placeholder_detected": False,
    })

    def mock_build(p):
        from docx import Document
        docx_path = str(tmp_path / "book.docx")
        doc = Document()
        doc.add_paragraph("Índice")
        doc.add_paragraph("Capítulo 1: El nacimiento de Internet")
        doc.add_paragraph("palabra " * fallback_words)
        doc.save(docx_path)
        return {"docx_path": docx_path, "book_id": 1001, "language": "es",
                "chapter_count": 1, "image_count": 0}

    monkeypatch.setattr(runner, "build_book_docx", mock_build)
    monkeypatch.setattr(runner, "save_checkpoint", lambda *a, **k: "cp.json")
    return {"fallback_text": fallback_text, "fc_capture": fc_capture}


def test_status_completed_only_after_docx(monkeypatch, tmp_path: Path) -> None:
    """Si el DOCX no se genera (build_book_docx lanza), el runner nunca devuelve completed."""
    import importlib
    importlib.reload(runner)

    monkeypatch.setattr(runner, "init_db", lambda: None)
    monkeypatch.setattr(runner, "plan_execute", lambda p: {
        "title": "Test", "provider": "ollama", "model": "test",
        "chapters": [{"image_requirements": 0}],
    })
    monkeypatch.setattr(runner, "research_execute", lambda p: {
        "status": "PASS",
        "sources": [
            {"title": f"s{i}", "url": f"https://space-lair.test/{i}",
             "source_type": "web", "relevance": "N/A",
             "accessed_at": "2024-01-01T00:00:00+00:00"}
            for i in range(3)
        ],
        "execution_mode": "real",
    })
    chapter_file = tmp_path / "chapter.md"
    chapter_file.write_text("palabra " * 1800, encoding="utf-8")
    real_open = open

    def mock_open(path, *args, **kwargs):
        if "chapter.md" in str(path):
            return real_open(chapter_file, *args, **kwargs)
        return real_open(path, *args, **kwargs)

    monkeypatch.setattr("builtins.open", mock_open)
    monkeypatch.setattr(runner, "chapter_execute", lambda p: {
        "word_count": 1800, "quality_gate": "PASS", "execution_mode": "real",
        "quality_errors": [], "chapter_md_path": str(chapter_file),
    })
    monkeypatch.setattr(runner, "_chapter_detect_placeholder", lambda t: False)
    monkeypatch.setattr(runner, "fc_execute", lambda p: {
        "status": "PASS", "quality_gate": "PASS", "claims_checked": 1,
        "supported_claims": 0, "unsupported_claims": 0, "conflicting_claims": 0,
    })
    monkeypatch.setattr(runner, "editor_execute", lambda p: {
        "status": "PASS", "quality_gate": "PASS", "input_words": 1800,
        "output_words": 1800, "execution_mode": "real", "placeholder_detected": False,
    })

    def boom_build(p):
        raise RuntimeError("simulando fallo en DOCX")

    monkeypatch.setattr(runner, "build_book_docx", boom_build)
    monkeypatch.setattr(runner, "save_checkpoint", lambda *a, **k: "cp.json")

    result = runner.run()
    assert result["status"] == "error"
    assert result["failed_stage"] == "document_builder"
    assert result["status"] != "completed"


def test_chapter_text_matches_fact_check_input_after_fallback(monkeypatch, tmp_path: Path) -> None:
    """Chapter Writer FAIL corta el pipeline; no hay fallback ni Fact Check."""
    mocks = _fallback_setup(monkeypatch, tmp_path)

    result = runner.run()

    assert result.get("fallback_used") is not True
    assert result["status"] == "error"
    assert result["failed_stage"] == "chapter"
    # Fact Check no se ejecuta porque el pipeline corta en chapter
    assert mocks["fc_capture"] == []


def test_book_planner_fail_aborts(monkeypatch, tmp_path: Path) -> None:
    """Book Planner sin capítulos → status error, failed_stage book_planner y no continúa."""
    import importlib
    importlib.reload(runner)

    research_called: list = []
    monkeypatch.setattr(runner, "init_db", lambda: None)
    monkeypatch.setattr(runner, "plan_execute", lambda p: {
        "title": "Test", "provider": "ollama", "model": "test",
        "chapters": [],  # plan inválido: sin capítulos
    })
    monkeypatch.setattr(runner, "research_execute", lambda p: research_called.append(1) or {})
    monkeypatch.setattr(runner, "save_checkpoint", lambda *a, **k: "cp.json")

    result = runner.run()

    assert result["status"] == "error"
    assert result["failed_stage"] == "book_planner"
    # Research/Outline/Chapter NO deben haberse ejecutado
    assert research_called == []


def test_book_planner_checkpoint_stage(monkeypatch, tmp_path: Path) -> None:
    """El checkpoint del Book Planner NO se registra como OUTLINE sino como BOOK_PLAN."""
    import importlib
    importlib.reload(runner)

    stages: list = []
    monkeypatch.setattr(runner, "init_db", lambda: None)
    monkeypatch.setattr(runner, "plan_execute", lambda p: {
        "title": "Test", "provider": "ollama", "model": "test",
        "chapters": [{"image_requirements": 0}],
    })
    monkeypatch.setattr(runner, "research_execute", lambda p: {
        "status": "PASS",
        "sources": [
            {"title": f"s{i}", "url": f"https://space-lair.test/{i}",
             "source_type": "web", "relevance": "N/A",
             "accessed_at": "2024-01-01T00:00:00+00:00"}
            for i in range(3)
        ],
        "execution_mode": "real",
    })
    chapter_file = tmp_path / "chapter.md"
    chapter_file.write_text("palabra " * 1800, encoding="utf-8")
    real_open = open

    def mock_open(path, *args, **kwargs):
        if "chapter.md" in str(path):
            return real_open(chapter_file, *args, **kwargs)
        return real_open(path, *args, **kwargs)

    monkeypatch.setattr("builtins.open", mock_open)
    monkeypatch.setattr(runner, "chapter_execute", lambda p: {
        "word_count": 1800, "quality_gate": "PASS", "execution_mode": "real",
        "quality_errors": [], "chapter_md_path": str(chapter_file),
    })
    monkeypatch.setattr(runner, "_chapter_detect_placeholder", lambda t: False)
    monkeypatch.setattr(runner, "fc_execute", lambda p: {
        "status": "PASS", "quality_gate": "PASS", "claims_checked": 1,
        "supported_claims": 0, "unsupported_claims": 0, "conflicting_claims": 0,
    })
    monkeypatch.setattr(runner, "editor_execute", lambda p: {
        "status": "PASS", "quality_gate": "PASS", "input_words": 1800,
        "output_words": 1800, "execution_mode": "real", "placeholder_detected": False,
    })

    def spy_save(book_id, stage, payload, **kw):
        stages.append(stage)
        return "cp.json"

    monkeypatch.setattr(runner, "save_checkpoint", spy_save)

    def mock_build(p):
        from docx import Document
        docx_path = str(tmp_path / "book.docx")
        doc = Document()
        doc.add_paragraph("Índice")
        doc.add_paragraph("Capítulo 1: El nacimiento de Internet")
        doc.add_paragraph("palabra " * 1800)
        doc.save(docx_path)
        return {"docx_path": docx_path, "book_id": 1001, "language": "es",
                "chapter_count": 1, "image_count": 0}

    monkeypatch.setattr(runner, "build_book_docx", mock_build)

    result = runner.run()

        # El primer checkpoint (Book Planner) usa BOOK_PLAN, no OUTLINE
    assert stages and stages[0] == "book_plan"
    # El outline real sí está presente más adelante
    assert "outline" in stages
    assert result["status"] == "completed"


def test_fallback_used_never_completes(monkeypatch, tmp_path: Path) -> None:
    """Invariante de integridad: chapter FAIL corta el pipeline; nunca usa fallback.

    El Chapter Writer falla por longitud y el runner corta en chapter.
    """
    _fallback_setup(monkeypatch, tmp_path)
    result = runner.run()
    assert result.get("fallback_used") is not True
    assert result["status"] == "error"
    assert result["failed_stage"] == "chapter"
    assert result["chapter_generation_status"] == "FAIL"


def test_fallback_final_qc_records_integrity_fields(monkeypatch, tmp_path: Path) -> None:
    """Chapter FAIL corta el pipeline; no se produce checkpoint FINAL_QC.

    El runner ya no usa fallback ni llega a QC cuando chapter_execute falla.
    Verificamos que el resultado sea error en chapter y que no exista
    checkpoint FINAL_QC.
    """
    captured: list = []
    _fallback_setup(monkeypatch, tmp_path)
    monkeypatch.setattr(
        runner, "save_checkpoint",
        lambda *a, **k: captured.append((a, k)) or "cp.json",
    )

    result = runner.run()

    assert result.get("fallback_used") is not True
    assert result["status"] == "error"
    assert result["failed_stage"] == "chapter"
    assert result["chapter_generation_status"] == "FAIL"
    # No se debe haber guardado un checkpoint FINAL_QC porque el pipeline cortó antes.
    final_qc_calls = [c for c in captured if c[0][1] == runner.Stage.FINAL_QC.value]
    assert final_qc_calls == []

