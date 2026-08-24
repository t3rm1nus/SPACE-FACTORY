"""Generación de libros DOCX profesionales a partir del modelo editorial."""

from __future__ import annotations

import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Mm, Pt, RGBColor

from core.book.book_schema import Book, Chapter
from core.schemas import BookDocxOutput, BookDocxPayload

logger = logging.getLogger(__name__)

DEFAULT_PAGE_SIZE = "A4"
DEFAULT_MARGINS_MM = {
    "top": 25.4,
    "bottom": 25.4,
    "left": 25.4,
    "right": 25.4,
}

# ---------------------------------------------------------------------------
# Tabla de strings de interfaz del DOCX (FASE 1 — multi-idioma)
#
# Todas las cadenas de texto visible en el documento (portada, legal, TOC,
# introducción, labels, captions, fuentes) se parametizan por idioma a través
# de este diccionario. El fallback es "es" (comportamiento histórico); un
# idioma desconocido nunca causa error — se usa el español como seguro.
#
# Los valores con placeholders ({year}, {author}, {title}, {n}, {index}) se
# formatean en el sitio con str.format() usando el idioma activo.
# ---------------------------------------------------------------------------
_UI_STRINGS = {
    "es": {
        "untitled": "Sin título",
        "title_label": "Título: {title}",
        "author_label": "Autor: {author}",
        "copyright_template": "© {year} {holder}. Todos los derechos reservados.",
        "reproduction_notice": (
            "Queda prohibida la reproducción total o parcial de esta obra, "
            "por cualquier medio o procedimiento, sin permiso expreso del autor."
        ),
        "toc_title": "Índice",
        "chapter_label": "Capítulo",
        "chapter_prefix": "Capítulo {n}: {title}",
        "chapter_fallback": "Capítulo {n}",
        "introduction_title": "Introducción",
        "no_content": "(Sin contenido disponible para este capítulo)",
        "sources_heading": "Fuentes utilizadas",
        "figure_prefix": "Figura {index}{ref}",
        "sources_attribution": "(fuente: {url})",
    },
    "en": {
        "untitled": "Untitled",
        "title_label": "Title: {title}",
        "title_label_no_colon": "Title",
        "author_label": "Author: {author}",
        "copyright_template": "© {year} {holder}. All rights reserved.",
        "reproduction_notice": (
            "Reproduction of this work in any form or by any means, "
            "without the express permission of the copyright holder, is prohibited."
        ),
        "toc_title": "Table of Contents",
        "chapter_label": "Chapter",
        "chapter_prefix": "Chapter {n}: {title}",
        "chapter_fallback": "Chapter {n}",
        "introduction_title": "Introduction",
        "no_content": "(No content available for this chapter)",
        "sources_heading": "Sources Used",
        "figure_prefix": "Figure {index}{ref}",
    },
}


def _ui(lang: str) -> dict:
    """Devuelve los strings de UI para el idioma solicitado, con fallback seguro a 'es'."""
    return _UI_STRINGS.get(lang, _UI_STRINGS["es"])
PAGE_SIZES_MM = {
    "A4": (210.0, 297.0),
    "LETTER": (215.9, 279.4),
    "LEGAL": (215.9, 355.6),
}

# ---------------------------------------------------------------------------
# Presets de maquetación (FASE 6)
# ---------------------------------------------------------------------------
# Valores por defecto por preset. Los overrides del usuario se aplican encima
# del preset seleccionado (ver `_effective_layout`).
_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}

PRESETS = {
    "editorial": {
        "font_family": "Georgia",
        "heading_font": "Georgia",
        "heading_color": "#1F3A5F",
        "body_alignment": "justify",
        "font_size": 11,
        "line_spacing": 1.15,
        "images_per_chapter": 3,
    },
    "moderno": {
        "font_family": "Arial",
        "heading_font": "Arial",
        "heading_color": "#6A3FB5",
        "body_alignment": "left",
        "font_size": 11,
        "line_spacing": 1.2,
        "images_per_chapter": 3,
    },
    "clasico": {
        "font_family": "Times New Roman",
        "heading_font": "Times New Roman",
        "heading_color": "#000000",
        "body_alignment": "justify",
        "font_size": 12,
        "line_spacing": 1.5,
        "images_per_chapter": 1,
    },
    "academico": {
        "font_family": "Garamond",
        "heading_font": "Garamond",
        "heading_color": "#1F3A5F",
        "body_alignment": "justify",
        "font_size": 11,
        "line_spacing": 1.5,
        "images_per_chapter": 1,
    },
    "dossier": {
        "font_family": "Arial",
        "heading_font": "Arial",
        "heading_color": "#000000",
        "body_alignment": "left",
        "font_size": 10,
        "line_spacing": 1.15,
        "images_per_chapter": 0,
    },
}

# Alias tolerantes (con/sin acentos) -> clave canónica
_PRESET_ALIASES = {
    "editorial": "editorial",
    "moderno": "moderno",
    "modern": "moderno",
    "clasico": "clasico",
    "classic": "clasico",
    "academico": "academico",
    "academic": "academico",
    "dossier": "dossier",
}

_HEADING_STYLES = [
    "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6",
]


def _parse_hex_color(value: Any) -> RGBColor:
    """Convierte un valor hexadecimal (#RRGGBB) o RGBColor a RGBColor."""
    if isinstance(value, RGBColor):
        return value
    if isinstance(value, str):
        s = value.strip().lstrip("#")
        if len(s) == 6:
            try:
                return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))
            except ValueError:
                pass
    return RGBColor(0x00, 0x00, 0x00)


def _effective_layout(layout_config: Optional[dict]) -> dict:
    """Resuelve la config efectiva: valores del preset + overrides del usuario.

    Devuelve un dict con claves normalizadas de `PRESETS`.
    """
    cfg = layout_config or {}
    preset_key = _PRESET_ALIASES.get(
        str(cfg.get("preset") or "editorial").strip().lower(), "editorial"
    )
    effective = dict(PRESETS[preset_key])

    overrides = cfg.get("overrides") or {}
    # Solo aplicamos overrides cuyas claves existen en el preset (ignora ruido)
    for key in ("font_family", "heading_font", "heading_color", "body_alignment",
                "font_size", "line_spacing", "images_per_chapter"):
        if key in overrides and overrides[key] not in (None, ""):
            effective[key] = overrides[key]

    # Normalización de tipos
    try:
        effective["font_size"] = Pt(float(effective["font_size"]))
    except (TypeError, ValueError):
        effective["font_size"] = Pt(11)
    try:
        effective["line_spacing"] = float(effective["line_spacing"])
    except (TypeError, ValueError):
        effective["line_spacing"] = 1.15
    align = _ALIGN_MAP.get(
        str(effective["body_alignment"]).strip().lower(), WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    effective["body_alignment"] = align
    try:
        effective["images_per_chapter"] = int(effective["images_per_chapter"])
    except (TypeError, ValueError):
        effective["images_per_chapter"] = 3
    return effective


def _apply_layout_config(doc: Document, layout_config: Optional[dict]) -> dict:
    """Aplica la config de maquetación a los estilos existentes del documento.

    Debe llamarse DESPUÉS de crear los estilos base (tras `_ensure_style`).
    Devuelve la config efectiva resuelta.
    """
    cfg = _effective_layout(layout_config)

    # Cuerpo: fuente, tamaño, alineación e interlineado
    for sname in ("Normal", "Intro Body", "TOC Entry", "Caption"):
        try:
            st = doc.styles[sname]
            st.font.name = cfg["font_family"]
            st.font.size = cfg["font_size"]
            st.paragraph_format.alignment = cfg["body_alignment"]
            if sname != "Caption":
                st.paragraph_format.line_spacing = cfg["line_spacing"]
        except KeyError:
            continue

    # Títulos: color y fuente
    heading_color = _parse_hex_color(cfg["heading_color"])
    heading_font = cfg.get("heading_font") or cfg["font_family"]
    for sname in _HEADING_STYLES:
        try:
            st = doc.styles[sname]
            st.font.color.rgb = heading_color
            st.font.name = heading_font
        except KeyError:
            continue

    return cfg


def layout_presets() -> dict:
    """Expone los presets en formato JSON seguro (colores como hex).

    Usado por el endpoint /api/layout-presets para poblar los selectores del Front.
    """
    return dict(PRESETS)


def _ensure_style(
    doc: Document,
    name: str,
    *,
    base_style: str = "Normal",
    font_name: str = "Calibri",
    font_size: Pt = Pt(11),
    bold: bool = False,
    italic: bool = False,
    alignment: Optional[int] = None,
    space_after: Pt = Pt(6),
    space_before: Pt = Pt(0),
    color: RGBColor = RGBColor(0x00, 0x00, 0x00),
) -> Any:
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, 1)
        style.base_style = doc.styles[base_style]

    font = style.font
    font.name = font_name
    font.size = font_size
    font.bold = bold
    font.italic = italic
    font.color.rgb = color

    pf = style.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    if alignment is not None:
        pf.alignment = alignment

    return style


def _apply_page_config(section: Any, page_config: Optional[dict]) -> None:
    size_name = (
        page_config.get("size", DEFAULT_PAGE_SIZE) if page_config else DEFAULT_PAGE_SIZE
    )
    width_mm = page_config.get("width_mm") if page_config else None
    height_mm = page_config.get("height_mm") if page_config else None
    margins = (
        page_config.get("margins_mm", DEFAULT_MARGINS_MM)
        if page_config
        else DEFAULT_MARGINS_MM
    )

    if width_mm and height_mm:
        section.page_width = Mm(float(width_mm))
        section.page_height = Mm(float(height_mm))
    else:
        w, h = PAGE_SIZES_MM.get(size_name.upper(), PAGE_SIZES_MM["A4"])
        section.page_width = Mm(w)
        section.page_height = Mm(h)

    section.top_margin = Mm(float(margins.get("top", DEFAULT_MARGINS_MM["top"])))
    section.bottom_margin = Mm(float(margins.get("bottom", DEFAULT_MARGINS_MM["bottom"])))
    section.left_margin = Mm(float(margins.get("left", DEFAULT_MARGINS_MM["left"])))
    section.right_margin = Mm(float(margins.get("right", DEFAULT_MARGINS_MM["right"])))


def _add_page_number(paragraph: Any) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def _add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    """Añade un hipervínculo REAL (<w:hyperlink>) a un párrafo.

    Mismo patrón de manipulación XML de bajo nivel que _add_page_number:
    relación externa en el part + elemento w:hyperlink con r:id. El run
    interior lleva subrayado y el color azul estándar de hyperlink.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    r_id = paragraph.part.relate_to(
        url, RT.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(color)
    rPr.append(u)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)\s]+)\)")


def _add_formatted_text(paragraph: Any, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|\[[^\]]+\]\([^)\s]+\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Enlaces markdown [texto](url) → hipervínculo real.
            m = _LINK_RE.fullmatch(part.strip())
            if m:
                _add_hyperlink(paragraph, m.group(1), m.group(2))
            else:
                paragraph.add_run(part)


def _iter_markdown_blocks(lines, *, skip_first_h1: bool = False):
    """Genera los bloques de markdown de un capítulo como tuplas (tipo, nivel, valor).

    Replica exactamente la clasificación de _parse_markdown_to_paragraphs:
    - ("heading", nivel, texto)
    - ("quote", None, texto)
    - ("separator", None, None)
    - ("bullet", None, texto)
    - ("normal", None, texto)

    Las líneas vacías se omiten. Si skip_first_h1 es True se salta el primer
    "# " del bloque (el titular ya se añadió explícitamente como Heading 1).
    """
    for i, line in enumerate(lines):
        stripped = line.rstrip()
        if not stripped.strip():
            continue

        # Saltar el primer # Heading si ya se añadió el título explícitamente
        if skip_first_h1 and i == 0 and stripped.startswith("# "):
            continue

        if stripped.startswith("###### "):
            yield ("heading", 6, stripped[7:].strip())
        elif stripped.startswith("##### "):
            yield ("heading", 5, stripped[6:].strip())
        elif stripped.startswith("#### "):
            yield ("heading", 4, stripped[5:].strip())
        elif stripped.startswith("### "):
            yield ("heading", 3, stripped[4:].strip())
        elif stripped.startswith("## "):
            yield ("heading", 2, stripped[3:].strip())
        elif stripped.startswith("# "):
            yield ("heading", 1, stripped[2:].strip())
        elif stripped.startswith("> "):
            yield ("quote", None, stripped[2:].strip())
        elif stripped.strip() == "---":
            yield ("separator", None, None)
        elif stripped.strip().startswith("- ") or stripped.strip().startswith("* "):
            yield ("bullet", None, stripped.strip()[2:])
        else:
            yield ("normal", None, stripped)


def _render_markdown_block(
    doc: Document, kind: str, level: Optional[int], value: Optional[str]
) -> None:
    """Renderiza un bloque de markdown (emitido por _iter_markdown_blocks) como párrafo."""
    if kind == "heading":
        doc.add_paragraph(value or "", style=f"Heading {int(level)}")
    elif kind == "quote":
        doc.add_paragraph(value or "", style="Quote")
    elif kind == "separator":
        p = doc.add_paragraph("— — —", style="Separator")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif kind == "bullet":
        p = doc.add_paragraph(style="List Bullet")
        _add_formatted_text(p, value or "")
    else:  # normal
        p = doc.add_paragraph(style="Normal")
        _add_formatted_text(p, value or "")


def _parse_markdown_to_paragraphs(doc: Document, text: str, *, skip_first_h1: bool = False) -> None:
    for kind, level, value in _iter_markdown_blocks(text.splitlines(), skip_first_h1=skip_first_h1):
        _render_markdown_block(doc, kind, level, value)


def _split_sources_tail(lines) -> tuple[list[str], list[str]]:
    """Divide las líneas del capítulo en cuerpo y cola de fuentes generada por el LLM.

    Detecta headings H2 tipo '## Fuentes utilizadas', '## Referencias' o
    '## Bibliografía'. La cola se DESCARTA (fix #9/#13/#14): la sección de
    fuentes del DOCX se construye de forma determinista desde chapter.sources
    (fuentes reales del payload), nunca desde texto fabricado por el LLM.
    Devuelve (body, tail); tail se conserva por compatibilidad de firma.
    """
    body: list[str] = []
    tail: list[str] = []
    seen: set[str] = set()
    in_sources = False
    for line in lines:
        stripped = line.rstrip()
        lowered = stripped.lower()
        if (
            not in_sources
            and stripped.startswith("## ")
            and (
                "fuentes utilizadas" in lowered
                or "referencias" in lowered
                or "bibliograf" in lowered
                or lowered.strip("# ").strip() == "fuentes"
                or "sources used" in lowered
                or "sources" in lowered
                or "references" in lowered
                or "bibliography" in lowered
            )
        ):
            in_sources = True
        if in_sources:
            # Mitigación defensiva (NO fix de causa raíz): deduplicar líneas
            # de fuente exactas (mismo texto tras normalizar espacios),
            # conservando la primera aparición y el orden. Las vacías se
            # conservan tal cual (separadores visuales).
            if stripped.strip():
                key = " ".join(stripped.split())
                if key in seen:
                    continue
                seen.add(key)
            tail.append(line)
        else:
            body.append(line)
    return body, tail


def _count_body_paragraphs(lines) -> int:
    """Cuenta los párrafos de contenido real que producirá el cuerpo del capítulo."""
    return sum(1 for _ in _iter_markdown_blocks(lines, skip_first_h1=True))


def _compute_insertion_points(num_images: int, num_paragraphs: int):
    """Calcula los puntos de inserción (~uniformes) de las imágenes entre párrafos.

    Devuelve:
    - [] si no hay imágenes (no cambia el comportamiento).
    - None si hay MENOS párrafos que imágenes (capítulo muy corto): señal para
      conservar el comportamiento actual (todas las imágenes al final).
    - lista ordenada de índices de párrafo (1-based) de contenido tras los
      cuales se coloca cada imagen, usando floor(k * M / (N + 1)) para k=1..N.

    Los puntos se fuerzan a >= 1 y sin duplicados para no colocar dos imágenes
    en el mismo hueco ni antes del primer párrafo.
    """
    if num_images <= 0:
        return []
    if num_paragraphs <= num_images:
        # Capítulo tan corto que no hay huecos suficientes: no repartir.
        return None
    points: list[int] = []
    for k in range(1, num_images + 1):
        p = (k * num_paragraphs) // (num_images + 1)
        p = max(1, p)
        if p not in points:
            points.append(p)
    return sorted(points)


def _add_body_with_images(doc: Document, body_lines, images, emit_image) -> None:
    """Añade el cuerpo del capítulo intercalando las imágenes entre los párrafos.

    ``emit_image(path)`` coloca una imagen (con su caption) justo después del
    párrafo corriente. Consume las rutas de ``images`` (pop) conforme las coloca,
    de modo que el llamador solo emite al final las que queden sin intercalar.
    """
    body = [ln for ln in body_lines]
    points = _compute_insertion_points(len(images), _count_body_paragraphs(body))
    if points is None:
        # Capítulo muy corto: conserva el comportamiento actual (todo al final).
        _parse_markdown_to_paragraphs(doc, "\n".join(body), skip_first_h1=True)
        return

    points_set = set(points)
    par_count = 0
    for kind, level, value in _iter_markdown_blocks(body, skip_first_h1=True):
        _render_markdown_block(doc, kind, level, value)
        par_count += 1
        if par_count in points_set and images:
            emit_image(images.pop(0))


def _prepare_image_source(image_path: str):
    """Prepara la fuente para ``doc.add_picture``.

    python-docx no soporta ``.webp`` nativamente (lanza excepción). Si el archivo
    es webp, se convierte a PNG en memoria con Pillow y se devuelve un
    ``BytesIO``; en cualquier otro caso se devuelve la ruta tal cual. Devuelve
    ``None`` si la ruta no existe o si la conversión webp falla (el caller
    mantiene el comportamiento de skip con warning, sin romper la fase).
    """
    if not image_path or not os.path.isfile(image_path):
        return None
    ext = os.path.splitext(image_path)[1].lower() if image_path else ""
    if ext != ".webp":
        return image_path
    try:
        import io

        from PIL import Image as PILImage

        img = PILImage.open(image_path)
        buf = io.BytesIO()
        # PNG no soporta modos con paleta/alpha tal cual en todos los casos; se
        # normaliza a RGBA/RGB según transparencia para conservar calidad.
        if img.mode in ("RGBA", "LA") or "transparency" in img.info:
            img = img.convert("RGBA")
        else:
            img = img.convert("RGB")
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf
    except Exception as exc:  # noqa: BLE001 - conversión fallida: skip con warning
        logger.warning("No se pudo convertir la imagen webp a PNG %s: %s", image_path, exc)
        return None


def _add_image_if_exists(
    doc: Document, image_path: str, caption_ref: str, index: int,
    language: str = "es",
) -> bool:
    """Inserta la imagen con su caption; devuelve True si se insertó de verdad."""
    if not image_path or not os.path.isfile(image_path):
        logger.warning("Imagen no encontrada, se omite: %s", image_path)
        return False

    image_source = _prepare_image_source(image_path)
    if image_source is None:
        logger.warning("Imagen no insertable (webp sin conversión), se omite: %s", image_path)
        return False

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(image_source, width=Inches(5.5))
    except Exception as exc:
        logger.warning("No se pudo insertar la imagen %s: %s", image_path, exc)
        return False

    if caption_ref is not None:
        ui = _ui(language)
        cap = doc.add_paragraph(
            ui["figure_prefix"].format(index=index, ref=caption_ref),
            style="Caption",
        )
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return True


def _add_cover(doc: Document, book: Book, language: str = "es") -> None:
    ui = _ui(language)
    doc.add_paragraph(book.title or ui["untitled"], style="Cover Title")
    if book.subtitle:
        doc.add_paragraph(book.subtitle, style="Cover Subtitle")
    if book.author:
        doc.add_paragraph(book.author, style="Cover Author")
    doc.add_page_break()


def _add_legal(doc: Document, book: Book, language: str = "es") -> None:
    # Año del copyright: se prefiere el año de creación del libro (semántica
    # editorial correcta: la obra se publica cuando se creó) y, si falta,
    # el año actual real (antes había un 2024 hardcodeado).
    ui = _ui(language)
    year = book.created_at.year if book.created_at else datetime.now().year
    author = book.author
    title = book.title or ui["untitled"]
    # Si no hay autor, se omite la línea de autor (igual que _add_cover) y
    # el copyright usa el título del libro en vez del autor.
    legal_lines = [ui["title_label"].format(title=title)]
    if author:
        legal_lines.append(ui["author_label"].format(author=author))
    copyright_holder = author or title
    legal_text = (
        "\n".join(legal_lines)
        + f"\n{ui['copyright_template'].format(year=year, holder=copyright_holder)}\n"
        f"{ui['reproduction_notice']}"
    )
    p = doc.add_paragraph(legal_text, style="Legal Text")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_page_break()


def _add_toc(doc: Document, chapters: list[Chapter], language: str = "es") -> None:
    ui = _ui(language)
    doc.add_paragraph(ui["toc_title"], style="TOC Title")
    # Fix defensivo: si el título ya trae prefijo "Capítulo N" / "Chapter N"
    # con o sin ":" (p.ej. baked-in por un plan antiguo, o plano "Chapter N"),
    # NO volver a prefijar para evitar duplicación.
    label = re.escape(ui["chapter_label"])
    chapter_prefix_re = re.compile(rf"(?i)^{label}\s+\d+\b")
    for ch in chapters:
        fallback_title = ui["chapter_fallback"].format(n=ch.number)
        title = ch.title or fallback_title
        if chapter_prefix_re.match(title):
            entry_text = title
        else:
            entry_text = ui["chapter_prefix"].format(n=ch.number, title=title)
        entry = doc.add_paragraph(entry_text, style="TOC Entry")
        entry.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT
        )
        dots = "." * max(1, 80 - len(entry_text))
        entry.add_run(f" {dots}")
    doc.add_page_break()


def _add_introduction(doc: Document, book: Book, language: str = "es") -> None:
    ui = _ui(language)
    if book.description:
        doc.add_paragraph(ui["introduction_title"], style="Intro Title")
        _parse_markdown_to_paragraphs(doc, book.description)
        # Sin page break aquí: _add_chapter lo añade al empezar


def _add_chapter(doc: Document, chapter: Chapter, language: str) -> None:
    doc.add_page_break()
    ui = _ui(language)
    title = chapter.title or ui["chapter_fallback"].format(n=chapter.number)
    doc.add_paragraph(title, style="Heading 1")

    content = getattr(chapter, f"edited_{language}") or getattr(chapter, f"draft_{language}")
    caption_ref = f": {title}"
    caption_number = 0
    images = list(chapter.images or [])

    def _emit_image(img_path: str) -> None:
        nonlocal caption_number
        # El número SOLO se consume cuando la imagen se inserta de verdad
        # (fix hueco: antes se incrementaba aunque el fichero no existiera,
        # dejando huecos tipo "Figura 2" sin "Figura 1").
        if _add_image_if_exists(doc, img_path, caption_ref, caption_number + 1, language):
            caption_number += 1

    if content:
        body, _sources_tail_discarded = _split_sources_tail(content.splitlines())
        # Fix #9/#13/#14: la cola de fuentes generada por el LLM (posible
        # duplicada, truncada o fabricada) se DESCARTA. La sección de fuentes
        # del DOCX se reconstruye de forma determinista desde chapter.sources
        # (URLs reales pobladas por _chapter_source_urls/SourceManager).
        _add_body_with_images(doc, body, images, _emit_image)
    else:
        doc.add_paragraph(
            ui["no_content"], style="Normal"
        )

    real_sources = [str(u).strip() for u in (chapter.sources or []) if str(u).strip()]
    if real_sources:
        doc.add_paragraph(ui["sources_heading"], style="Heading 2")
        for url in real_sources:
            src_para = doc.add_paragraph(style="List Bullet")
            try:
                _add_hyperlink(src_para, url, url)
            except Exception:
                # Nunca romper el DOCX por una URL malformada: texto plano.
                src_para.add_run(url)

    # Imágenes que quedaron sin intercalar (capítulo muy corto / huecos insuficientes):
    # se conserva el comportamiento actual, todas al final.
    for img_path in images:
        _emit_image(img_path)


def build_book_docx(payload: dict[str, Any]) -> dict[str, Any]:
    validated = BookDocxPayload(**payload)
    book = Book.model_validate(validated.book)
    language = validated.language
    page_config = validated.page_config or {}

    output_dir = Path("output") / "docx"
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"book_{book.book_id}_{language}.docx"
    docx_path = str(output_dir / filename)

    doc = Document()

    section = doc.sections[0]
    _apply_page_config(section, page_config)

    _ensure_style(doc, "Cover Title", font_size=Pt(28), bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    _ensure_style(doc, "Cover Subtitle", font_size=Pt(18),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))
    _ensure_style(doc, "Cover Author", font_size=Pt(14),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(36))
    _ensure_style(doc, "Legal Text", font_size=Pt(10),
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(12))
    _ensure_style(doc, "TOC Title", font_size=Pt(18), bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    _ensure_style(doc, "TOC Entry", font_size=Pt(11),
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6))
    _ensure_style(doc, "Intro Title", font_size=Pt(22), bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    _ensure_style(doc, "Intro Body", font_size=Pt(11),
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(12))
    _ensure_style(doc, "Caption", font_size=Pt(10), italic=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    _ensure_style(doc, "Heading 4", font_size=Pt(12), bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6), space_before=Pt(12),
                  color=RGBColor(0x1F, 0x3A, 0x5F))
    _ensure_style(doc, "Heading 5", font_size=Pt(11), bold=True, italic=True,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6), space_before=Pt(10),
                  color=RGBColor(0x1F, 0x3A, 0x5F))
    _ensure_style(doc, "Heading 6", font_size=Pt(10), bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4), space_before=Pt(8),
                  color=RGBColor(0x1F, 0x3A, 0x5F))
    _ensure_style(doc, "Quote", font_size=Pt(11), italic=True,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(8),
                  color=RGBColor(0x44, 0x44, 0x44))
    _ensure_style(doc, "Separator", font_size=Pt(9),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4),
                  color=RGBColor(0x88, 0x88, 0x88))

    doc.core_properties.title = book.title or ""
    doc.core_properties.author = book.author or ""
    doc.core_properties.subject = book.genre or ""
    doc.core_properties.comments = (book.description or "")[:255]
    doc.core_properties.language = language

    # La PORTADA (primera página) no lleva header ni footer (estilo editorial
    # estándar). La página LEGAL (página 2) SÍ hereda header/footer normales.
    section.different_first_page_header_footer = True

    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.text = book.title or ""
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Solo número de página: nunca exponer el nombre de archivo interno al lector
    # (el título ya aparece en el header de cada página).
    _add_page_number(footer_para)

    chapters = sorted(book.chapters or [], key=lambda c: c.number)

    _add_cover(doc, book, language)
    _add_legal(doc, book, language)
    _add_toc(doc, chapters, language)
    _add_introduction(doc, book, language)
    for ch in chapters:
        _add_chapter(doc, ch, language)

    # Aplicar preset/overrides de maquetación sobre los estilos ya creados (FASE 6)
    _apply_layout_config(doc, getattr(book, "layout_config", None))

    doc.save(docx_path)

    return BookDocxOutput(
        docx_path=docx_path,
        book_id=book.book_id or 0,
        language=language,
        chapter_count=len(chapters),
        image_count=sum(len(ch.images) for ch in chapters),
    ).model_dump()


def health_check() -> dict[str, Any]:
    try:
        import docx  # noqa: F401
        return {"healthy": True, "dependencies": {"python-docx": "ok"}}
    except Exception as exc:
        return {"healthy": False, "error": str(exc), "dependencies": {"python-docx": "error"}}

def execute(payload: dict, capability: str = "build_book_docx") -> dict:
    """Wrapper de ejecución exigido por el registro de módulos.

    Delega en build_book_docx(payload) para generar el DOCX del libro.
    """
    return build_book_docx(payload)

