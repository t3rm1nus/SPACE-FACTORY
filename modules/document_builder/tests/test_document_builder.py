"""Tests del módulo document_builder."""

from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path
from unittest.mock import patch

import pytest
from PIL import Image as PILImage

from core.book.book_schema import Book, Chapter
from core.schemas import BookDocxPayload
from modules.document_builder.main import build_book_docx, health_check


@pytest.fixture(autouse=True)
def _isolate_output(tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    yield


def test_health_check():
    result = health_check()
    assert result["healthy"] is True
    assert result["dependencies"]["python-docx"] == "ok"


def _book_payload(tmp_path: Path, with_images: bool = True) -> dict:
    img_path = tmp_path / "hero.png"
    if with_images:
        PILImage.new("RGB", (64, 64), color="red").save(img_path)

    book = {
        "book_id": 1,
        "title": "El libro del espacio",
        "subtitle": "Manual de diseño profesional",
        "description": "Esta es la introducción del libro.",
        "author": "Space Lair",
        "target_audience": "desarrolladores",
        "genre": "tecnología",
        "languages": ["es"],
        "target_chapters": 2,
        "status": "edited",
        "created_at": datetime(2024, 1, 1).isoformat(),
        "chapters": [
                    {
                        "chapter_id": 10,
                        "book_id": 1,
                        "number": 1,
                        "title": "Introducción",
                        "edited_es": "## Antecedentes\n\nContenido del capítulo 1.",
                        "images": [str(img_path)] if with_images else [],
                    },
                    {
                        "chapter_id": 11,
                        "book_id": 1,
                        "number": 2,
                        "title": "Diseño avanzado",
                        "draft_es": "## Técnicas\n\nContenido del capítulo 2.",
                        "images": [],
                    },
                ],
    }
    return {
        "book": book,
        "language": "es",
        "page_config": {
            "size": "A4",
            "margins_mm": {"top": 20, "bottom": 20, "left": 20, "right": 20},
        },
    }


def test_build_book_docx_creates_real_docx(tmp_path: Path):
    payload = _book_payload(tmp_path, with_images=True)
    out = build_book_docx(payload)

    assert os.path.isfile(out["docx_path"])
    assert out["language"] == "es"
    assert out["chapter_count"] == 2
    assert out["image_count"] == 1
    assert out["docx_path"].endswith(f"book_{out['book_id']}_es.docx")

    from docx import Document
    doc = Document(out["docx_path"])

    assert doc.core_properties.title == "El libro del espacio"
    assert doc.core_properties.author == "Space Lair"
    assert doc.core_properties.language == "es"

    texts = [p.text for p in doc.paragraphs]
    assert "El libro del espacio" in texts
    assert "El libro del espacio" in texts  # cover
    assert any("Introducción" in t for t in texts)
    assert any("Diseño avanzado" in t for t in texts)
    assert any("Contenido del capítulo 1" in t for t in texts)
    assert any("Contenido del capítulo 2" in t for t in texts)
    assert any("Figura 1" in t for t in texts)

    # Check images embedded
    assert len(doc.inline_shapes) == 1

    # Check header/footer
    section = doc.sections[0]
    assert section.different_first_page_header_footer is True
    assert "El libro del espacio" in section.header.paragraphs[0].text

    # A1: año legal = año de creación del libro (2024 en el fixture),
    # nunca un literal hardcodeado desincronizado.
    legal_texts = [p.text for p in doc.paragraphs if "©" in p.text]
    assert any("© 2024" in t for t in legal_texts)

    # A2: la portada (primera página) no lleva header/footer.
    first_page_header_text = "".join(
        p.text for p in section.first_page_header.paragraphs
    )
    first_page_footer_text = "".join(
        p.text for p in section.first_page_footer.paragraphs
    )
    assert "El libro del espacio" not in first_page_header_text

    # A3: el footer NO expone el nombre de archivo interno; solo número de
    # página (campo PAGE, texto vacío hasta renderizarse).
    footer_text = section.footer.paragraphs[0].text
    footer_xml = section.footer.paragraphs[0]._p.xml
    assert ".docx" not in footer_text
    assert f"book_{out['book_id']}" not in footer_text
    assert 'w:instrText' in footer_xml and "PAGE" in footer_xml


def test_build_book_docx_custom_page_size(tmp_path: Path):
    payload = _book_payload(tmp_path, with_images=False)
    payload["page_config"] = {"size": "LETTER", "margins_mm": {"top": 15, "bottom": 15, "left": 15, "right": 15}}
    out = build_book_docx(payload)

    from docx import Document
    doc = Document(out["docx_path"])
    section = doc.sections[0]
    assert round(section.page_width.inches, 2) == 8.5
    assert round(section.page_height.inches, 2) == 11.0
    assert round(section.top_margin.mm) == 15


def test_build_book_docx_missing_images_are_skipped(tmp_path: Path):
    payload = _book_payload(tmp_path, with_images=False)
    # Force an image path that does not exist
    payload["book"]["chapters"][0]["images"] = [str(tmp_path / "no_existe.png")]
    out = build_book_docx(payload)

    assert os.path.isfile(out["docx_path"])  # should not crash
    from docx import Document
    doc = Document(out["docx_path"])  # should open fine
    assert len(doc.inline_shapes) == 0


def test_build_book_docx_no_duplicate_chapters(tmp_path: Path):
    payload = _book_payload(tmp_path, with_images=False)
    # chapters are already unique in payload
    out = build_book_docx(payload)
    assert out["chapter_count"] == 2

    from docx import Document
    doc = Document(out["docx_path"])
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert headings.count("Introducción") == 1
    assert headings.count("Diseño avanzado") == 1


def test_build_book_docx_uses_edited_over_draft(tmp_path: Path):
    payload = _book_payload(tmp_path, with_images=False)
    payload["book"]["chapters"][0]["edited_es"] = "## Editado\n\nContenido editado."
    out = build_book_docx(payload)

    from docx import Document
    doc = Document(out["docx_path"])  # should not crash
    texts = [p.text for p in doc.paragraphs]
    assert "Contenido editado." in texts
    assert "Contenido del capítulo 1" not in texts


# ---------------------------------------------------------------------------
# 8E.6C — ISOLAMIENTO MULTI-LIBRO POR book_id
# Dos libros del mismo idioma NO deben compartir ruta ni sobreescribirse.
# ---------------------------------------------------------------------------
def test_docx_isolation_two_books_same_language(tmp_path: Path):
    """8E.6C: 1 book_id = 1 artefacto DOCX independiente.

    Secuencia obligada: generar A, inspeccionar, generar B, comprobar que A
    no cambió (ni bytes ni identidad) y que B es B. El cross-overwrite entre
    libros del mismo idioma está eliminado estructuralmente.
    """
    import hashlib
    from docx import Document

    def _payload(book_id: int, title: str) -> dict:
        book = {
            "book_id": book_id,
            "title": title,
            "subtitle": "Manual de algo",
            "description": "Descripción.",
            "author": "Space Lair",
            "target_audience": "General",
            "genre": "Divulgación",
            "languages": ["es"],
            "target_chapters": 1,
            "status": "edited",
            "created_at": datetime(2024, 1, 1).isoformat(),
            "chapters": [
                {
                    "chapter_id": book_id * 10,
                    "book_id": book_id,
                    "number": 1,
                    "title": "Capítulo 1",
                    "edited_es": "## Técnicas\n\nContenido del capítulo.",
                    "images": [],
                },
            ],
        }
        return {
            "book": book,
            "language": "es",
            "page_config": {"size": "A4", "margins_mm": {"top": 25, "bottom": 25, "left": 25, "right": 25}},
        }

    # A (español)
    out_a = build_book_docx(_payload(11, "Libro A"))
    path_a = out_a["docx_path"]
    with open(path_a, "rb") as fh:
        bytes_a_before = fh.read()
    hash_a_before = hashlib.sha256(bytes_a_before).hexdigest()
    title_a_before = Document(path_a).core_properties.title

    # B (español) generado DESPUÉS de A
    out_b = build_book_docx(_payload(22, "Libro B"))
    path_b = out_b["docx_path"]
    with open(path_b, "rb") as fh:
        bytes_b = fh.read()
    title_b = Document(path_b).core_properties.title

    # A sigue siendo A (no fue sobreescrito por B)
    with open(path_a, "rb") as fh:
        hash_a_after = hashlib.sha256(fh.read()).hexdigest()
    title_a_after = Document(path_a).core_properties.title

    # 1) paths distintos
    assert path_a != path_b
    # 2) nombres canónicos por book_id
    assert path_a.endswith("book_11_es.docx")
    assert path_b.endswith("book_22_es.docx")
    # 3) ambos existen físicamente
    assert os.path.isfile(path_a)
    assert os.path.isfile(path_b)
    # 4) A no cambió al generar B (isolated)
    assert hash_a_before == hash_a_after
    assert title_a_before == title_a_after == "Libro A"
    # 5) B contiene su propia identidad
    assert title_b == "Libro B"
    assert bytes_a_before != bytes_b
    # 6) identidad 1:1: el path canónico de B es book_22_es.docx y el footer
    # NO expone nombres de archivo internos (A3: footer solo número de página).
    assert path_b.endswith("book_22_es.docx")
    footer_b = Document(path_b).sections[0].footer.paragraphs[0].text
    assert ".docx" not in footer_b


@pytest.mark.parametrize("preset,exp_font,exp_color", [
    ("moderno", "Arial", (0x6A, 0x3F, 0xB5)),
    ("editorial", "Georgia", (0x1F, 0x3A, 0x5F)),
    ("clasico", "Times New Roman", (0x00, 0x00, 0x00)),
])
def test_build_book_docx_applies_layout_preset(tmp_path, preset, exp_font, exp_color):
    """FASE 6: el preset de maquetación se aplica a los estilos del DOCX."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    payload = _book_payload(tmp_path, with_images=False)
    payload["book"]["layout_config"] = {"preset": preset, "overrides": {}}
    out = build_book_docx(payload)

    doc = Document(out["docx_path"])
    normal = doc.styles["Normal"]
    heading1 = doc.styles["Heading 1"]

    assert normal.font.name == exp_font
    rgb = heading1.font.color.rgb
    assert (rgb[0], rgb[1], rgb[2]) == exp_color


def test_build_book_docx_layout_overrides_win(tmp_path):
    """FASE 6: los overrides manuales anulan los valores del preset."""
    from docx import Document

    payload = _book_payload(tmp_path, with_images=False)
    payload["book"]["layout_config"] = {
        "preset": "editorial",
        "overrides": {
            "font_family": "Courier New",
            "heading_color": "#000000",
            "body_alignment": "left",
        },
    }
    out = build_book_docx(payload)

    doc = Document(out["docx_path"])
    assert doc.styles["Normal"].font.name == "Courier New"
    rgb = doc.styles["Heading 1"].font.color.rgb
    assert (rgb[0], rgb[1], rgb[2]) == (0x00, 0x00, 0x00)


def test_build_book_docx_legal_without_author_omits_line_and_uses_title(tmp_path):
    """Cuando book.author es None, la página legal NO muestra 'Autor: Autor'
    y el copyright usa el título del libro en vez del autor."""
    from docx import Document

    payload = _book_payload(tmp_path, with_images=False)
    payload["book"]["author"] = None
    out = build_book_docx(payload)

    doc = Document(out["docx_path"])
    texts = [p.text for p in doc.paragraphs]

    # No debe aparecer el placeholder "Autor: Autor"
    assert not any("Autor: Autor" in t for t in texts)
    # La línea del copyright usa el título del libro
    assert any("© 2024 El libro del espacio. Todos los derechos reservados." in t for t in texts)

# ---------------------------------------------------------------------------
# 7.9D.8 — REPARTO GEOMÉTRICO DE IMÁGENES ENTRE PÁRRAFOS
# Antes de la fase 7.9D.8 todas las imágenes de un capítulo se insertaban
# juntas al final (tras "Fuentes utilizadas"). Ahora se intercalan de forma
# aproximadamente uniforme entre los párrafos de contenido real.
# ---------------------------------------------------------------------------
def _make_image_files(tmp_path: Path, count: int) -> list[str]:
    """Crea `count` imágenes PNG válidas y devuelve sus rutas."""
    paths = []
    for i in range(1, count + 1):
        p = tmp_path / f"fig_{i}.png"
        PILImage.new("RGB", (64, 64), color=(30, (i * 40) % 255, 120)).save(p)
        paths.append(str(p))
    return paths


def _single_chapter_payload(tmp_path: Path, edited_es: str, images_count: int) -> dict:
    """Payload con un único capítulo cuyo edited_es e imágenes son controlables."""
    book = {
        "book_id": 1,
        "title": "Libro reparto de imágenes",
        "subtitle": "Subtítulo",
        "description": "Introducción.",
        "author": "Space Lair",
        "target_audience": "General",
        "genre": "Divulgación",
        "languages": ["es"],
        "target_chapters": 1,
        "status": "edited",
        "created_at": datetime(2024, 1, 1).isoformat(),
        "chapters": [
            {
                "chapter_id": 10,
                "book_id": 1,
                "number": 1,
                "title": "Capítulo de reparto",
                "edited_es": edited_es,
                "images": _make_image_files(tmp_path, images_count),
            },
        ],
    }
    return {"book": book, "language": "es"}


def _image_paragraph_indices(doc) -> list[int]:
    """Índices (dentro de doc.paragraphs) de los párrafos que contienen un drawing."""
    indices = []
    for i, p in enumerate(doc.paragraphs):
        if p._p.xpath(".//w:drawing"):
            indices.append(i)
    return indices


def test_images_distributed_between_paragraphs(tmp_path):
    """Capítulo con 20 párrafos y 3 imágenes: no deben quedar consecutivas al final.

    Debe haber al menos un párrafo de texto real entre imagen1-imagen2 y
    entre imagen2-imagen3.
    """
    from docx import Document

    paras = "\n\n".join(f"Párrafo de contenido número {i}." for i in range(1, 21))
    payload = _single_chapter_payload(tmp_path, paras, images_count=3)
    out = build_book_docx(payload)

    doc = Document(out["docx_path"])
    draw = _image_paragraph_indices(doc)
    assert len(draw) == 3

    # Están en orden creciente y separadas (no consecutivas).
    assert draw == sorted(draw)
    assert all(d > 0 for d in draw)

    for prev, cur in zip(draw, draw[1:]):
        between = doc.paragraphs[prev + 1:cur]
        assert any(p.text.startswith("Párrafo de contenido") for p in between)

    # Hay texto real después de la última imagen (no queda pegada al final).
    assert any(p.text.startswith("Párrafo de contenido")
               for p in doc.paragraphs[draw[-1] + 1:])

def test_images_short_chapter_all_at_end(tmp_path):
    """Capítulo muy corto (2 párrafos) con 3 imágenes: cae en la salvaguarda.

    No debe romper ni intentar repartir en huecos que no existen: las imágenes
    van todas al final, tras el último párrafo de texto.
    """
    from docx import Document

    payload = _single_chapter_payload(tmp_path, "Párrafo A.\n\nPárrafo B.", images_count=3)
    out = build_book_docx(payload)

    doc = Document(out["docx_path"])
    draw = _image_paragraph_indices(doc)
    assert len(draw) == 3

    last_text_idx = max(
        i for i, p in enumerate(doc.paragraphs)
        if p.text.startswith("Párrafo")
    )
    # Todas las imágenes van después del último párrafo de texto real.
    assert all(i > last_text_idx for i in draw)

    texts = [p.text for p in doc.paragraphs]
    assert any(t.startswith("Figura 1") for t in texts)
    assert any(t.startswith("Figura 3") for t in texts)


def test_fuentes_utilizadas_after_last_image(tmp_path):
    """'## Fuentes utilizadas' (construida desde chapter.sources) sigue al final
    del capítulo, tras las imágenes intercaladas. Fix #9/#13/#14."""
    from docx import Document

    paras = "\n\n".join(f"Párrafo {i}." for i in range(1, 21))
    content = (
        paras
        + "\n\n## Fuentes utilizadas\n\nReferencia uno.\n\nReferencia dos."
    )
    payload = _single_chapter_payload(tmp_path, content, images_count=3)
    # Fuentes reales del payload (vía _chapter_source_urls/SourceManager).
    payload["book"]["chapters"][0]["sources"] = [
        "https://es.wikipedia.org/wiki/Fuente_uno",
        "https://es.wikipedia.org/wiki/Fuente_dos",
    ]
    out = build_book_docx(payload)

    doc = Document(out["docx_path"])
    draw = _image_paragraph_indices(doc)
    assert len(draw) == 3

    texts = [p.text for p in doc.paragraphs]
    idx_sources = next(i for i, p in enumerate(doc.paragraphs) if p.text == "Fuentes utilizadas")
    idx_ref1 = next(i for i, t in enumerate(texts) if t.endswith("/Fuente_uno"))
    idx_ref2 = next(i for i, t in enumerate(texts) if t.endswith("/Fuente_dos"))

    # Todas las imágenes intercaladas van ANTES del bloque de fuentes.
    assert all(i < idx_sources for i in draw)
    # "Fuentes utilizadas" precede a sus referencias reales.
    assert idx_sources < idx_ref1 < idx_ref2
    # La cola fabricada por el LLM NO se renderiza.
    assert "Referencia uno." not in texts
    assert "Referencia dos." not in texts
    # Entre imágenes sigue habiendo texto real (reparto intacto).
    for prev, cur in zip(draw, draw[1:]):
        assert any(p.text.startswith("Párrafo") for p in doc.paragraphs[prev + 1:cur])


def test_no_images_behavior_unchanged(tmp_path):
    """Regresión: un capítulo sin imágenes NO cambia de comportamiento.

    No se generan drawings ni captions 'Figura N' y todo el texto está presente.
    """
    from docx import Document

    paras = "\n\n".join(f"Párrafo {i}." for i in range(1, 21))
    payload = _single_chapter_payload(tmp_path, paras, images_count=0)
    out = build_book_docx(payload)

    doc = Document(out["docx_path"])
    assert _image_paragraph_indices(doc) == []

    texts = [p.text for p in doc.paragraphs]
    assert not any(t.startswith("Figura") for t in texts)
    assert sum(1 for t in texts if t.startswith("Párrafo")) == 20


def test_markdown_link_rendered_as_real_hyperlink(tmp_path):
    """B1: [Título](url) en el CUERPO se renderiza como hipervínculo real.

    (Fix #9/#13/#14: la cola '## Fuentes utilizadas' del LLM ya no se renderiza;
    este test cubre el render de enlaces markdown dentro del cuerpo.)
    """
    from docx import Document

    content = (
        "Párrafo introductorio con [Título](https://ejemplo.com/pagina) en el cuerpo."
    )
    payload = _single_chapter_payload(tmp_path, content, images_count=0)
    out = build_book_docx(payload)

    doc = Document(out["docx_path"])
    texts = [p.text for p in doc.paragraphs]
    full_text = "\n".join(texts)
    assert "](http" not in full_text
    assert "[Título]" not in full_text

    # Hipervínculo real: elemento w:hyperlink en el XML + relación externa.
    xml = "\n".join(p._p.xml for p in doc.paragraphs)
    assert "<w:hyperlink" in xml
    rels = doc.part.rels
    assert any(
        r.reltype.endswith("/hyperlink") and r.target_ref == "https://ejemplo.com/pagina"
        for r in rels.values()
    )


def test_duplicate_source_lines_discarded(tmp_path):
    """B2 (fix #9): la cola duplicada del LLM NO aparece en el DOCX; solo las
    fuentes reales de chapter.sources."""
    from docx import Document

    src = "Fuente repetida de ejemplo (web_searxng)"
    content = (
        "Párrafo uno.\n\n"
        f"## Fuentes utilizadas\n\n- {src}\n\n- Otra fuente distinta (web_wikipedia)\n\n- {src}"
    )
    payload = _single_chapter_payload(tmp_path, content, images_count=0)
    payload["book"]["chapters"][0]["sources"] = [
        "https://es.wikipedia.org/wiki/Fuente_real",
    ]
    out = build_book_docx(payload)

    texts = [p.text for p in Document(out["docx_path"]).paragraphs]
    assert texts.count(src) == 0  # descartada, no solo deduplicada
    assert "Otra fuente distinta" not in "".join(texts)
    assert any(t.endswith("/Fuente_real") for t in texts)


def test_sources_section_from_chapter_sources_without_llm_tail(tmp_path):
    """Fix #14: aunque draft_es/edited_es NO traiga ninguna cola de fuentes,
    la sección 'Fuentes utilizadas' se genera desde chapter.sources."""
    from docx import Document

    payload = _single_chapter_payload(tmp_path, "Párrafo único sin cola.", images_count=0)
    payload["book"]["chapters"][0]["sources"] = [
        "https://es.wikipedia.org/wiki/A",
        "https://www.abacus.coop/guia",
        "https://www.amazon.es/dp/TEST",
    ]
    out = build_book_docx(payload)

    doc = Document(out["docx_path"])
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    texts = [p.text for p in doc.paragraphs]

    assert "Fuentes utilizadas" in headings
    for url_part in ("/A", "/guia", "/dp/TEST"):
        assert any(t.endswith(url_part) for t in texts)


def test_empty_chapter_sources_omits_section_no_placeholder(tmp_path):
    """Fix #14/#13: sin chapter.sources NO se inventa ni placeholder: se omite
    la sección completa, incluso si el LLM dejó una cola en el texto."""
    from docx import Document

    content = (
        "Párrafo uno.\n\n"
        "## Fuentes utilizadas\n\n- Smith & Johnson (2019). Fake Editorial."
    )
    payload = _single_chapter_payload(tmp_path, content, images_count=0)
    payload["book"]["chapters"][0]["sources"] = []
    out = build_book_docx(payload)

    texts = [p.text for p in Document(out["docx_path"]).paragraphs]
    assert "Fuentes utilizadas" not in texts
    assert not any("Smith & Johnson" in t or "Fake Editorial" in t for t in texts)
    assert "(Sin fuentes proporcionadas)" not in texts


def test_fabricated_llm_tail_replaced_by_real_sources(tmp_path):
    """Fix #13: si el LLM fabricó citas APA falsas, NO aparecen en el DOCX;
    solo las reales de chapter.sources."""
    from docx import Document

    content = (
        "Párrafo uno.\n\n"
        "## Fuentes\n\n"
        "- Smith, J. & Johnson, K. (2019). Trust Theory. Fake Press. ISBN 000-0.\n"
        "- Brown, A. (2020). Confidence. White & Black Books."
    )
    payload = _single_chapter_payload(tmp_path, content, images_count=0)
    payload["book"]["chapters"][0]["sources"] = [
        "https://es.wikipedia.org/wiki/Confianza",
    ]
    out = build_book_docx(payload)

    full_text = "\n".join(p.text for p in Document(out["docx_path"]).paragraphs)
    assert "Smith" not in full_text and "Johnson" not in full_text
    assert "Brown" not in full_text and "ISBN" not in full_text
    assert any(t.endswith("/Confianza") for t in full_text.splitlines())


def test_figure_numbering_no_gap_when_first_image_missing(tmp_path):
    """B3: si la primera imagen no existe, la segunda es 'Figura 1' (sin hueco)."""
    from docx import Document

    payload = _single_chapter_payload(tmp_path, "Párrafo A.", images_count=1)
    missing = str(tmp_path / "no_existe.png")
    existing = payload["book"]["chapters"][0]["images"][0]
    payload["book"]["chapters"][0]["images"] = [missing, existing]

    out = build_book_docx(payload)

    texts = [p.text for p in Document(out["docx_path"]).paragraphs]
    captions = [t for t in texts if t.startswith("Figura")]
    assert len(captions) == 1
    assert captions[0].startswith("Figura 1")


def _toc_chapter(number: int, title: str) -> Chapter:
    return Chapter(
        chapter_id=number,
        book_id=1,
        number=number,
        title=title,
    )


def _toc_doc(chapters):
    """Document con el estilo 'TOC Entry' asegurado (como hace build_book_docx)."""
    from docx import Document
    from docx.shared import Pt

    from modules.document_builder.main import _add_toc, _ensure_style

    doc = Document()
    _ensure_style(doc, "TOC Title", font_size=Pt(18), bold=True)
    _ensure_style(doc, "TOC Entry", font_size=Pt(11), alignment=__import__("docx").enum.text.WD_ALIGN_PARAGRAPH.LEFT)
    _add_toc(doc, chapters)
    return doc


def test_add_toc_does_not_duplicate_prefix():
    """Título que ya trae 'Capítulo N:' NO debe duplicar el prefijo en el TOC."""
    chapters = [_toc_chapter(3, "Capítulo 3: Algo ya prefijado")]
    doc = _toc_doc(chapters)
    entries = [p.text for p in doc.paragraphs if "Algo ya prefijado" in p.text]
    assert len(entries) == 1
    # Una sola aparición del prefijo (no "Capítulo 3: Capítulo 3:")
    assert entries[0].count("Capítulo 3:") == 1


def test_add_toc_prefixes_plain_title():
    """Caso sano: título sin prefijo SÍ recibe 'Capítulo N: ' normalmente."""
    chapters = [_toc_chapter(2, "El Origen del Rock")]
    doc = _toc_doc(chapters)
    entries = [p.text for p in doc.paragraphs if "El Origen del Rock" in p.text]
    assert len(entries) == 1
    assert entries[0].startswith("Capítulo 2: El Origen del Rock")


def test_webp_image_is_inserted(tmp_path):
    """§17 #18 — una imagen .webp SÍ debe insertarse en el DOCX (se convierte a
    PNG en memoria en _prepare_image_source), en vez de saltarse con warning por
    no ser soportada por python-docx."""
    from docx import Document

    from modules.document_builder.main import _prepare_image_source

    webp = tmp_path / "test_web.webp"
    PILImage.new("RGB", (64, 64), color="blue").save(webp, format="WEBP")

    # Conversión en memoria: no webp -> devuelve la ruta; webp -> BytesIO PNG.
    assert _prepare_image_source(str(webp)) is not None

    # Integración: build_book_docx con una imagen webp produce la Figura.
    payload = _single_chapter_payload(tmp_path, "Párrafo A.", images_count=1)
    payload["book"]["chapters"][0]["images"] = [str(webp)]

    out = build_book_docx(payload)
    texts = [p.text for p in Document(out["docx_path"]).paragraphs]
    captions = [t for t in texts if t.startswith("Figura")]
    assert len(captions) == 1
    assert captions[0].startswith("Figura 1")


def test_add_toc_plain_capitulo_n_not_duplicated():
    """Caso books 49-54: título plano 'Capítulo N' (sin ':') NO debe duplicarse."""
    chapters = [_toc_chapter(1, "Capítulo 1")]
    doc = _toc_doc(chapters)
    # La entrada lleva puntos de relleno al final; verificamos el inicio del
    # texto y que NO se haya duplicado el prefijo.
    entries = [p.text for p in doc.paragraphs if p.text.startswith("Capítulo 1")]
    assert len(entries) == 1
    assert not entries[0].startswith("Capítulo 1: ")  # sin doble prefijo
# ---------------------------------------------------------------------------
# FASE 1 — PARAMETRIZACIÓN POR IDIOMA (feature "generación en inglés")
# Los strings de interfaz del DOCX (portada, legal, TOC, introducción, labels,
# captions, fuentes) se traducen según `language`, con "es" como fallback.
# ---------------------------------------------------------------------------
def _en_payload(tmp_path: Path, edited_en: str, images_count: int = 0) -> dict:
    """Payload con un único capítulo en inglés (edited_en) e imágenes."""
    book = {
        "book_id": 1,
        "title": "English Book Title",
        "subtitle": "Subtitle",
        "description": "Introduction.",
        "author": "Space Lair",
        "target_audience": "General",
        "genre": "Nonfiction",
        "languages": ["en"],
        "target_chapters": 1,
        "status": "edited",
        "created_at": datetime(2024, 1, 1).isoformat(),
        "chapters": [
            {
                "chapter_id": 10,
                "book_id": 1,
                "number": 1,
                "title": "Chapter One",
                "edited_en": edited_en,
                "images": _make_image_files(tmp_path, images_count) if images_count else [],
            },
        ],
    }
    return {"book": book, "language": "en"}


def test_es_ui_strings_unchanged(tmp_path):
    """FASE 1 regresión: con language='es' (default) los strings de UI son los
    mismos que antes de la parametrización (comportamiento específico)."""
    from docx import Document

    payload = _single_chapter_payload(
        tmp_path, "## Tema\n\nContenido.", images_count=0
    )
    payload["book"]["chapters"][0]["sources"] = [
        "https://es.wikipedia.org/wiki/Fuente_real"
    ]
    out = build_book_docx(payload)
    doc = Document(out["docx_path"])
    texts = [p.text for p in doc.paragraphs]

    # Los marcadores en español se conservan exactamente.
    assert "Índice" in texts
    assert "Introducción" in texts
    assert "Fuentes utilizadas" in texts
    assert any(t.startswith("Título: Libro reparto de imágenes") for t in texts)
    assert any("© 2024 Space Lair. Todos los derechos reservados." in t for t in texts)
    # No debe aparecer vocabulario en inglés.
    assert not any("Table of Contents" in t for t in texts)
    assert not any("Sources Used" in t for t in texts)


def test_en_ui_strings_translated(tmp_path):
    """FASE 1: con language='en' el DOCX usa los strings en inglés."""
    from docx import Document

    payload = _en_payload(tmp_path, "## Topic\n\nChapter content.", images_count=0)
    payload["book"]["chapters"][0]["sources"] = [
        "https://en.wikipedia.org/wiki/Real_source",
    ]
    out = build_book_docx(payload)
    doc = Document(out["docx_path"])
    texts = [p.text for p in doc.paragraphs]

    # Claves en inglés.
    assert "Table of Contents" in texts
    assert "Introduction" in texts
    assert "Sources Used" in texts
    assert any(t.startswith("Title: English Book Title") for t in texts)
    assert any("© 2024 Space Lair. All rights reserved." in t for t in texts)
    # No debe colarse el fallback español en estos marcadores clave.
    assert not any(t == "Índice" for t in texts)
    assert not any(t == "Fuentes utilizadas" for t in texts)
    assert not any("Todos los derechos reservados" in t for t in texts)


def test_en_figure_caption_translated(tmp_path):
    """FASE 1: el caption de imagen usa 'Figure' en inglés (no 'Figura')."""
    from docx import Document

    payload = _en_payload(tmp_path, "Párrafo A.\n\nPárrafo B.", images_count=1)
    out = build_book_docx(payload)
    texts = [p.text for p in Document(out["docx_path"]).paragraphs]
    captions = [t for t in texts if t.startswith("Figure")]
    assert len(captions) == 1
    assert captions[0].startswith("Figure 1")


def test_split_sources_tail_discards_english():
    """FASE 1: _split_sources_tail descarta una cola de fuentes en inglés
    ('## Sources Used' + lista) de la misma forma que la española."""
    from modules.document_builder.main import _split_sources_tail

    content = "\n".join([
        "Párrafo uno.",
        "",
        "## Sources used",
        "",
        "- https://en.wikipedia.org/wiki/A",
        "- https://en.wikipedia.org/wiki/B",
    ])
    body, tail = _split_sources_tail(content.splitlines())
    joined = "\n".join(body)
    assert "Sources used" not in joined
    assert "en.wikipedia.org" not in joined
    assert "Párrafo uno." in joined
    assert tail  # la cola se captura pero se descarta aguas abajo
